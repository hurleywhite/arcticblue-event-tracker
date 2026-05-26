#!/usr/bin/env python3
"""
extract-urls.py — pull every verified event URL out of the source .docx by
walking its hyperlink relationships.

NO HALLUCINATION RULE:
  Only URLs that exist as actual hyperlinks in the source document are kept.
  An event with no hyperlink stays unlinked in the output. The next AI
  agent reading this code: do not "fill in" URLs by guessing. If you want
  to add one, add it manually to data/event-urls-manual.json (see below).

Output: data/event-urls-from-doc.json — a {"event_number": "url"} map.
The build script merges this with data/event-urls-manual.json (if present),
with manual entries taking precedence.

Usage:
  python3 src/extract-urls.py
"""
import sys
sys.path.insert(0, '/Users/hurleywhite/Library/Python/3.11/lib/python/site-packages')
from docx import Document
import re
import json
from pathlib import Path

HERE = Path(__file__).resolve().parent.parent
DOC_PATH = HERE / 'data' / 'ArcticBlue AI 2026 Event Tracker.docx'
OUT_PATH = HERE / 'data' / 'event-urls-from-doc.json'

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

# Sub-page paths to demote — never invented, just preferred to skip if a
# cleaner root URL exists for the same event in the source doc.
DEMOTE = ('/speakers', '/sponsorship', '/sponsors', '/downloads', '/cfp', '/sponsorship-prospectus')


def iter_paragraphs(doc):
    """Yield every paragraph, including those inside table cells."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def hyperlinks_in(p, url_by_rid):
    """Return the list of URLs hyperlinked in this paragraph."""
    urls = []
    for hyper in p._p.findall(f'.//{W_NS}hyperlink'):
        rid = hyper.get(f'{R_NS}id')
        if rid and rid in url_by_rid:
            urls.append(url_by_rid[rid].rstrip('.,;:'))
    return urls


def pick_primary(urls):
    """Among verified URLs for an event, prefer the cleanest root."""
    if not urls:
        return None

    def score(u):
        low = u.lower()
        demoted = any(s in low for s in DEMOTE)
        return (1 if demoted else 0, len(u))

    return sorted(urls, key=score)[0]


def main():
    doc = Document(str(DOC_PATH))
    rels = doc.part.rels
    url_by_rid = {
        rid: rel.target_ref
        for rid, rel in rels.items()
        if 'hyperlink' in rel.reltype.lower()
    }

    events = []
    current = None
    for p in iter_paragraphs(doc):
        txt = p.text.strip()
        m = re.match(r'^(\d+)\.\s+(.+?)\s+—\s+(.+?)\s+\|\s+(.+)$', txt) if txt else None
        if m:
            if current is not None:
                events.append(current)
            current = {'num': int(m.group(1)), 'name': m.group(2), 'urls': []}
            # Capture hyperlinks on the event-header line itself
            for u in hyperlinks_in(p, url_by_rid):
                if u not in current['urls']:
                    current['urls'].append(u)
            continue
        if current is None:
            continue
        for u in hyperlinks_in(p, url_by_rid):
            if u not in current['urls']:
                current['urls'].append(u)

    if current is not None:
        events.append(current)

    mapping = {}
    for e in events:
        primary = pick_primary(e['urls'])
        if primary:
            mapping[str(e['num'])] = primary

    OUT_PATH.write_text(json.dumps(mapping, indent=2) + '\n', encoding='utf-8')
    print(f'Verified URLs extracted: {len(mapping)} of {len(events)} events')
    print(f'Wrote: {OUT_PATH}')


if __name__ == '__main__':
    main()
